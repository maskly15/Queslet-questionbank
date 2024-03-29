


#Upload File
from ..forms import UploadFileForm
from django.core.files.storage import default_storage
from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import redirect
from django.contrib.auth import login

from ..models.models import Mcq,Subject,SubjectAccess
from ..ultils import isManger
from ..api import api_encode
from docx import Document
import os
import shutil
import time
import re
import easyocr
from PIL import Image

import torch
import numpy as np
from sentence_transformers import SentenceTransformer
from sentence_transformers import SentenceTransformer, util
from ..Dbcontext import connector


#regex find image_path 
img_regex = "(?<=\[file: ).+?(?=\])"
ques_regex = "(\[file:).*?(\])"
op_regex = "([a-zA-Z]\.)"

# Easy OCR 
reader = easyocr.Reader(['en','vi'])

print("Connecting to pinecone")
conn = connector()

#Sbert 

def import_view(request):
   

    if request.method == "POST" :

        if 'submit_import' in request.POST:
            print("IMPORT NEW MCQ !")
            request.session['forward'] = False
            return submit_import(request)
        else:
            #Import
            form = UploadFileForm(request.POST,request.FILES)
            files = request.FILES.getlist('file')
            list_docx={}
            list_images={} 
            importSubject = request.session['subjectImport']
            importSubject = Subject.objects.get(subject=importSubject)
            #get Subject object
            #validate and report:
            #Load Images and docx file  
            for file in files:
                name_file,ex = os.path.splitext(str(file))
                # document = Document(file)
                if ex in ['.docx']:
                    print("found document")
                    list_docx[str(file)] = []
                    doc = Document(file)
                    for table in doc.tables:
                        
                        new_mcq = row2mcq(table,importSubject)
                        list_docx[str(file)].append(new_mcq)
                    # print("Total questions:",len(doc.tables))
                else:
                    list_images[str(file)] = file
                    print(str(file))
                    save_path = "temp/"+file.name
                    default_storage.save(save_path, file)
                    #reading Url 
                    # file_url = default_storage.url(save_path)
                    # print("Image_path:",file_url)
            duplicate_mcq =[]
            num_dup = 0
            num_dup_db = 0
            num_dup_doc = 0
            num_dup_mix = 0

            mcqs_set={}
            mcq_lst = []
            mcq_lst_encode = []
            not_found = []

            # check if have docx file 
            if len(list_docx) == 0:
                form = UploadFileForm()
                return render(request, 'import.html',{"mess":"Not found .docx file !",'form':form,'subjectImport':importSubject})

            for key in list_docx.keys():
                for mcq in list_docx[key]:
                    #store in session 
                    print(mcq.qid)
                    mcq_image_text =""
                    if mcq.contain_img:
                        try:
                            file_img = list_images[mcq.q_image]
                            mcq_image_text = ocr2Text(file_img)
                        except:
                            print("NOT FOUNd:",mcq.q_image)
                            mcq_image_text=""
                            not_found.append(mcq.q_image)
                    mcq_form = mcq.getMcq(mcq_image_text)
                        # print(mcq_form)
                    # encode = SbertModel.encode(str(mcq_form)).tolist()
                    encode = api_encode(str(mcq_form))
                    print(str(mcq.subject).strip())
                    result = conn.query_mcqs_encode(encode,str(mcq.subject).strip(),k=5)
                        # Filter dictionary by keeping elements whose keys are divisible by 2
                    list_id_dup = [ (d["id"], np.round(d["score"],2) ) for d in result["matches"] if d["score"] >= 0.7] 
                    print( list_id_dup)
                    mcqs_set[mcq.qid] = {"id":mcq.qid,"question":mcq.question,"image":mcq.q_image,"options":mcq.options,"answer":mcq.answer_q,"subject":str(mcq.subject),"haveImage":mcq.contain_img,"encode":encode,"qid":mcq.qid ,"dup_type":0  } 

                    if len(list_id_dup) != 0:
                        duplicate_mcq.append((mcq,list_id_dup,1))
                        num_dup += len(list_id_dup)
                        num_dup_db += len(list_id_dup)
                        mcqs_set[mcq.qid]["dup_type"] = 1
                    mcq_lst.append(mcq.qid)
                    mcq_lst_encode.append(encode)

            # compare all mcq in docs 
            pairs = cosin_pair(mcq_lst_encode)
            for pair in pairs:
                i, j = pair['index']
                temp_mcq = mcqs_set[mcq_lst[i]]
                # print("temp_mcq",temp_mcq)
                duplicate_mcq.append((temp_mcq,[(mcq_lst[j], np.round(pair['score'] ,2),2  )],2))
              
                num_dup += 1
                num_dup_doc += 1
                print("{} \t\t {} \t\t Score: {:.4f}".format(mcq_lst[i], mcq_lst[j], pair['score']))

                # kiem tra da bi trung trong database chua 

                if mcqs_set[mcq_lst[i]]['dup_type'] == 1:
                    mcqs_set[mcq_lst[i]]['dup_type'] = 3
                    num_dup_mix += 1
                    num_dup_doc -= 1
                    # num_dup_db -= 1

                else:
                    mcqs_set[mcq_lst[i]]['dup_type'] = 2

            request.session['temp_mcq'] = mcqs_set
            print("import length:",len(mcqs_set))
            # print("import length:",mcqs_set)     
            return render(request, 'import.html',{'dup_result':duplicate_mcq,"num_dup":num_dup,"num_dup_mix":num_dup_mix,"num_dup_doc":num_dup_doc,"num_dup_db":num_dup_db,"not_found":not_found})
    else:

         #authen the subject import 
        subjectImport = request.GET.get("subject")
        currentUser = request.user
        checkAuthen = authenUserSubject(subjectImport,currentUser)
        if checkAuthen:
            print("Ok !")
        else:
            return redirect('home')

        request.session['subjectImport'] =  subjectImport
        form = UploadFileForm()
        request.session['forward'] = True
    return render(request, 'import.html',{'form':form,'subjectImport':subjectImport})



def submit_import(request):
    checked_items = request.POST.getlist("check_import")
    #show import mcq
    print(checked_items)
    print("--------------------------")
    temp_dct = request.session['temp_mcq']
    currentUser = request.user
     
    for key_mcq in checked_items:
        mcq = temp_dct[key_mcq]
        encode = mcq["encode"]
        subject = mcq["subject"]


        #generate new id
        user_name = request.user.username
        new_id = subject+"-"+user_name+"-"+get_time()
        #get attr
        question = mcq["question"]
        options = mcq["options"]
        answer = mcq["answer"]
        haveImage = mcq["haveImage"]

        subject = Subject.objects.get(subject=subject)

        if haveImage:
            images = new_id+".png"
            temp_path = "media/temp/"+mcq["image"]
            image_path = "media/images/"+images
            new_path = "images/"+images
            shutil.move(temp_path, image_path)
            save_mcq= Mcq( 
                        qid =  new_id,
                        question = question,
                        options = options,
                        q_image = images,
                        answer_q = answer,
                        subject = subject,
                        contain_img = haveImage,
                        img_file = new_path ,
                        user = currentUser
                        )
                    # save_mcq.save()
        else:
            save_mcq= Mcq( 
                        qid =  new_id,
                        question = question,
                        options = options,
                        answer_q = answer,
                        subject = subject,
                        contain_img = haveImage,
                        user = currentUser
                         )
            images = ""
                    
        # save into database
        save_mcq.save()
        print("Import Mcq success !")
        print(save_mcq.subject)
        print(save_mcq.user)

        # save into Pinecone
        result = conn.upload(qid=new_id,encode=encode,question=question,contain=haveImage,q_image=images,subject=str(subject) )
        print(result)
            # Remove file in temp 
    shutil.rmtree('media/temp')
    User = request.user
    # request.session.flush()
    # login(User)
    return redirect('home')

def authenUserSubject(subject,Teacher):
    print("Import subject",subject)
    if subject == None:
        return False
    
    if isManger(Teacher):
        print("is manager !")
        return True
    try:
        ImportSubject = Subject.objects.get(subject=subject)
    except:
        print("Cant not find subject")
        return False
    
    try:
        print(Teacher.username)
        getAccess = SubjectAccess.objects.get(teacher = Teacher.username, subject = ImportSubject)
        return True
    except:
        print("Cant not find Access")
        return False 


def cosin_pair(encode):
  cosine_scores = util.cos_sim(encode, encode)
  pairs = []

  for i in range(len(cosine_scores)-1):
    for j in range(i+1, len(cosine_scores)):
        # Compare thresh hold 
        score = np.array(cosine_scores[i][j])
        if cosine_scores[i][j] >= 0.75:
            pairs.append({'index': [i, j], 'score': score})

  #Sort scores in decreasing order
  pairs = sorted(pairs, key=lambda x: x['score'], reverse=True)
  return pairs


def ocr2Text(file):
    path="temp.png"
    img = Image.open(file)
    img.save(path)
    img_text = reader.readtext(path, paragraph=True)

    question = ""
    for i in img_text:
        question += i[1] +" "
    return question

def get_time():
    t = time.time()
    t_s = int(t)
    return str(t_s)

def row2mcq(table,subject):
    data = [[cell.text for cell in row.cells] for row in table.rows]
    mcq_options = []
    ans =""
    for row in data:
        if(row[0].strip() == "answer"):
            ans = row[1]
        if("QN" in row[0]):
            mcq_name = row[0]
            text = row[1]
            text = text.replace("\n"," ")
            path = re.findall(img_regex,text) 
            if(len(path) == 0 ):
                path =""
            else:
                path = path[0]                
            question = re.sub(ques_regex, '',text )
            mcq_question = question
            # print("Question:",question)
            mcq_images = path
       
        if(re.match(op_regex,row[0])):
            text = row[1]
            text = text.replace("\n"," ")
            if len(text.strip()) != 0 :
                mcq_options.append(text)
    # print(mcq_images)
    # print(mcq_options)
    contain_img = False if len(mcq_images.strip()) <=0 else True
    if contain_img:
        new_mcq = Mcq(qid = mcq_name,question = mcq_question,options = str(mcq_options) , q_image = mcq_images , answer_q = ans, subject = subject, contain_img = contain_img, img_file = "temp/"+  mcq_images )
    else:
        new_mcq = Mcq(qid = mcq_name,question = mcq_question,options = str(mcq_options) , q_image = mcq_images , answer_q = ans, subject = subject, contain_img = contain_img )
    return new_mcq




